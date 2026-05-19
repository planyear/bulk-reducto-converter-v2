from pathlib import Path

from docx import Document


def convert(path: Path) -> str:
    doc = Document(str(path))
    blocks: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            blocks.append("\n".join(rows))

    return "\n\n".join(blocks)
